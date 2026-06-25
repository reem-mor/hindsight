#!/usr/bin/env python3
"""HINDSIGHT document extractor.

Called by the n8n "Execute Command" node. Given a file path, it extracts plain
text and any embedded images (so the workflow can optionally send dashboard
screenshots to Gemini Vision), and prints a single JSON object to stdout for n8n
to parse.

Usage:
    python extract_document.py /path/to/postmortem.pdf [--image-dir /tmp/hs_images]

Output (stdout, JSON):
    {
      "filename": "postmortem.pdf",
      "file_type": "pdf",
      "extracted_text": "....",
      "char_count": 1234,
      "images": [{"path": "/tmp/hs_images/postmortem_p1_0.png", "page": 1}],
      "ok": true
    }

Dependencies are imported lazily so the script degrades gracefully if a parser
is missing (it reports the problem in JSON instead of crashing the workflow).
"""
from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile


def _err(filename: str, file_type: str, message: str) -> dict:
    return {
        "filename": filename,
        "file_type": file_type,
        "extracted_text": "",
        "char_count": 0,
        "images": [],
        "ok": False,
        "error": message,
    }


def extract_pdf(path: str, image_dir: str) -> dict:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _err(os.path.basename(path), "pdf", "PyMuPDF (fitz) not installed: pip install pymupdf")

    text_parts: list[str] = []
    images: list[dict] = []
    base = os.path.splitext(os.path.basename(path))[0]

    try:
        with fitz.open(path) as doc:
            for page_index, page in enumerate(doc, start=1):
                text_parts.append(page.get_text("text"))
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    pix = None
                    try:
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha >= 4:  # CMYK/other → convert to RGB
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        out = os.path.join(image_dir, f"{base}_p{page_index}_{img_index}.png")
                        pix.save(out)
                        images.append({"path": out, "page": page_index})
                    except Exception as exc:  # noqa: BLE001
                        # Don't fail the whole extraction on one bad image.
                        images.append({"path": None, "page": page_index, "error": str(exc)})
                    finally:
                        pix = None  # release native pixmap memory promptly
    except Exception as exc:  # noqa: BLE001
        # Corrupt / password-protected / non-PDF content: degrade to a clean,
        # structured error instead of crashing the workflow node.
        return _err(os.path.basename(path), "pdf", f"could not parse PDF: {exc}")

    text = "\n".join(text_parts).strip()
    return {
        "filename": os.path.basename(path),
        "file_type": "pdf",
        "extracted_text": text,
        "char_count": len(text),
        "images": images,
        "ok": True,
    }


def extract_docx(path: str) -> dict:
    try:
        import docx  # python-docx
    except ImportError:
        return _err(os.path.basename(path), "docx", "python-docx not installed: pip install python-docx")

    try:
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs]
        # Include table cell text — postmortems often use tables for timelines.
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        # Corrupt / password-protected / non-OOXML content: degrade to JSON.
        return _err(os.path.basename(path), "docx", f"could not parse DOCX: {exc}")
    return {
        "filename": os.path.basename(path),
        "file_type": "docx",
        "extracted_text": text,
        "char_count": len(text),
        "images": [],
        "ok": True,
    }


def extract_text(path: str, file_type: str) -> dict:
    with open(path, "rb") as fh:
        raw = fh.read()
    # Try strict UTF-8 (BOM-tolerant) first, then Windows cp1252 (common for
    # exported .txt/.log on Windows), then a lossy UTF-8 read as last resort.
    text = None
    for enc in ("utf-8-sig", "cp1252"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="replace")
    text = text.strip()
    return {
        "filename": os.path.basename(path),
        "file_type": file_type,
        "extracted_text": text,
        "char_count": len(text),
        "images": [],
        "ok": True,
    }


IMAGE_EXTS = ("png", "jpg", "jpeg", "gif", "webp", "bmp", "tif", "tiff")


def extract_image(path: str, file_type: str) -> dict:
    """A standalone image (e.g. a SIEM dashboard screenshot) is Vision-only: it
    has no text layer, but the file itself is the image the Vision branch reads."""
    return {
        "filename": os.path.basename(path),
        "file_type": file_type,
        "extracted_text": "",
        "char_count": 0,
        "images": [{"path": path, "page": 0}],
        "ok": True,
    }


def extract(path: str, image_dir: str) -> dict:
    if not os.path.isfile(path):
        return _err(os.path.basename(path), "unknown", f"file not found: {path}")
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    if os.path.getsize(path) == 0:
        return _err(os.path.basename(path), ext, "empty file (0 bytes)")
    os.makedirs(image_dir, exist_ok=True)
    if ext == "pdf":
        result = extract_pdf(path, image_dir)
    elif ext == "docx":
        result = extract_docx(path)
    elif ext in ("txt", "md", "markdown", "log"):
        result = extract_text(path, ext)
    elif ext in IMAGE_EXTS:
        result = extract_image(path, ext)
    else:
        return _err(os.path.basename(path), ext, f"unsupported file type: {ext}")

    # Skip text-less documents UNLESS they carry a usable image for the Vision
    # branch (a scanned/image-only PDF still has something to analyse). Only count
    # images that actually exported — failed extractions are recorded as
    # {"path": null, ...} and must not be mistaken for Vision-ready content.
    usable_image = any(img.get("path") for img in result.get("images", []))
    if result.get("ok") and not result.get("extracted_text", "").strip() and not usable_image:
        return _err(
            result["filename"], result["file_type"],
            "no extractable text (empty or scanned document with no text layer or usable images)",
        )
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text + images from a document.")
    parser.add_argument("path", help="Path to the document")
    parser.add_argument(
        "--image-dir",
        default=os.path.join(tempfile.gettempdir(), "hindsight_images"),
        help="Directory to write extracted images",
    )
    args = parser.parse_args()
    try:
        result = extract(args.path, args.image_dir)
    except Exception as exc:  # noqa: BLE001
        # Last line of defence: the n8n node must always receive JSON, never a
        # Python traceback on stdout.
        result = _err(os.path.basename(args.path), "unknown", f"unexpected error: {exc}")
    payload = json.dumps(result, ensure_ascii=False)
    # Emit UTF-8 explicitly so non-ASCII text can't crash on a cp1252 console.
    sys.stdout.buffer.write((payload + "\n").encode("utf-8"))
    sys.stdout.buffer.flush()
    return 0 if result.get("ok") else 1


if __name__ == "__main__":
    sys.exit(main())
