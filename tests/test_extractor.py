"""Tests for document extractor (BON-1 Vision path).

The extractor always emits UTF-8 bytes on stdout, so every subprocess call
decodes with encoding="utf-8" (Windows' default cp1252 would mangle non-ASCII).
This mirrors how the n8n Execute Command node reads the stdout buffer.
"""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
EXTRACTOR = ROOT / "extractors" / "extract_document.py"
SAMPLES = ROOT / "samples"


def _run(*args: str, check: bool = False) -> dict:
    proc = subprocess.run(
        [sys.executable, str(EXTRACTOR), *args],
        capture_output=True,
        text=True,
        encoding="utf-8",
        check=check,
        cwd=str(ROOT),
    )
    return json.loads(proc.stdout)


def test_extractor_plain_markdown() -> None:
    data = _run(str(SAMPLES / "vuln_scan_critical_openssl.md"), check=True)
    assert data["ok"] is True
    assert data["file_type"] == "md"
    assert "CVSS" in data["extracted_text"]
    assert data["char_count"] > 100


def test_extractor_pdf_sample_with_image(tmp_path) -> None:
    """BON-1: a real PDF must extract text AND export at least one embedded image.

    Skips visibly (not a silent no-op) if the fixture is absent.
    """
    import pytest

    pdf = SAMPLES / "vuln_scan_sev1_critical_rce.pdf"
    if not pdf.is_file():
        pytest.skip(f"fixture missing: {pdf.name} (run samples/make_cyber_pdf_sample.py)")
    data = _run(str(pdf), "--image-dir", str(tmp_path), check=True)
    assert data["ok"] is True
    assert data["file_type"] == "pdf"
    assert data["char_count"] > 0
    assert "CVSS" in data["extracted_text"]
    # The Vision branch needs at least one embedded image actually exported.
    assert any(img.get("path") for img in data["images"])


def test_extractor_corrupt_docx_returns_structured_error(tmp_path) -> None:
    """A non-DOCX file with a .docx name must degrade to JSON, not crash the node."""
    bad = tmp_path / "corrupt.docx"
    bad.write_bytes(b"this is plainly not a real OOXML package")
    data = _run(str(bad))  # must be valid JSON even on failure
    assert data["ok"] is False
    assert data["file_type"] == "docx"
    assert "docx" in data["error"].lower()


def test_extractor_text_cp1252_fallback(tmp_path) -> None:
    """Windows-encoded text (cp1252 smart quote / en dash) must not become U+FFFD."""
    f = tmp_path / "win.txt"
    # 0x92 = right single quote (U+2019), 0x96 = en dash (U+2013) in cp1252.
    f.write_bytes(b"Owner\x92s incident note \x96 patch the edge gateway now")
    data = _run(str(f))
    text = data["extracted_text"]
    assert data["ok"] is True
    assert chr(0xFFFD) not in text  # no replacement chars
    assert chr(0x2019) in text  # 0x92 decoded to right single quote
    assert chr(0x2013) in text  # 0x96 decoded to en dash


def test_extractor_standalone_image_vision_ready(tmp_path) -> None:
    """A standalone image (e.g. a SIEM screenshot) is Vision-ready: no text, image present."""
    png = tmp_path / "siem_dashboard.png"
    # minimal valid 1x1 PNG
    png.write_bytes(
        bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
            "0000000a49444154789c6360000002000100052a0d0a2db40000000049454e44ae426082"
        )
    )
    data = _run(str(png))
    assert data["ok"] is True
    assert data["file_type"] == "png"
    assert data["extracted_text"] == ""
    assert any(img.get("path") for img in data["images"])  # the file itself -> Vision


def test_extractor_docx_happy_path(tmp_path) -> None:
    """DOCX: paragraphs AND table cells must be extracted (postmortems use tables)."""
    try:
        import docx as docx_mod
    except ImportError:
        import pytest

        pytest.skip("python-docx not installed")
    path = tmp_path / "incident.docx"
    d = docx_mod.Document()
    d.add_paragraph("Intrusion summary: brute-force against the bastion host.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CVE"
    table.rows[0].cells[1].text = "CVE-2026-9001"
    d.save(str(path))
    data = _run(str(path))
    assert data["ok"] is True
    assert data["file_type"] == "docx"
    assert "brute-force" in data["extracted_text"]
    assert "CVE-2026-9001" in data["extracted_text"]  # table cell text


def test_extractor_plain_txt_happy_path(tmp_path) -> None:
    """TXT: plain UTF-8 log/export extracts cleanly with the .txt file_type."""
    path = tmp_path / "alert.txt"
    path.write_text("SIEM alert: 412 failed SSH logins from 203.0.113.5 in 5 minutes.", encoding="utf-8")
    data = _run(str(path))
    assert data["ok"] is True
    assert data["file_type"] == "txt"
    assert "failed SSH logins" in data["extracted_text"]


def test_extractor_empty_docx_rejected(tmp_path) -> None:
    docx = tmp_path / "empty.docx"
    try:
        import docx as docx_mod
    except ImportError:
        return
    docx_mod.Document().save(str(docx))
    data = _run(str(docx))
    assert data["ok"] is False
    assert "no extractable text" in data["error"]
